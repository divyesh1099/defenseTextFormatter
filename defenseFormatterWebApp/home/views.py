import json
from django.shortcuts import redirect, render
from django.http import HttpResponse, JsonResponse, HttpResponseBadRequest
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
import google.generativeai as genai 
from googleApiKey import googleApiKey  
from docx import Document
import os
from django.conf import settings
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.pdfgen import canvas
import shutil

# Configure the generative AI model
genai.configure(api_key=googleApiKey)
model = genai.GenerativeModel('gemini-pro')
chat = model.start_chat()

def preprocessAI(raw_text):
    """
    This function constructs a detailed prompt for the AI to format the text
    according to specific rules and sends it through the chat interface.
    """
    instructions = '''
        "Analyze the document provided and generate a structured outline. For each section, include a marker (e.g., [TITLE], [HEADING], [LIST], [CODE]) followed by a placeholder text (e.g., 'Placeholder for section content'). Ensure the outline is concise, preserving the original document's structure without including any specific content details."
        Document Content:    
    '''
    # Combine instructions with the raw text
    full_message = f"{instructions}\n\n{raw_text}"
    
    # Send the combined message to the AI chat model
    response = chat.send_message(full_message)
    
    # Assume response.text contains the AI-formatted text
    formatted_text = response.text 

    return formatted_text
    # return "This is a test response without calling the external API."

# Function to apply formatting based on section type
def apply_formatting(document, section_type, content):
    if section_type == '[TITLE]':
        p = document.add_paragraph()
        run = p.add_run(content)
        run.bold = True
        run.font.size = Pt(24)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif section_type == '[HEADING]':
        p = document.add_paragraph()
        run = p.add_run(content)
        run.bold = True
        run.font.size = Pt(16)
    elif section_type == '[LIST]':
        document.add_paragraph(content, style='ListBullet')
    elif section_type == '[CODE]':
        p = document.add_paragraph()
        run = p.add_run(content)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    else:  # For normal text
        document.add_paragraph(content)

# Function to process the AI-generated outline and create a document
def create_document_from_outline(ai_output):
    document = Document()
    sections = re.split(r'(\[TITLE\]|\[HEADING\]|\[LIST\]|\[CODE\])', ai_output)
    
    section_type = None
    for i, section in enumerate(sections):
        if section in ['[TITLE]', '[HEADING]', '[LIST]', '[CODE]']:
            section_type = section
        elif section_type:
            apply_formatting(document, section_type, section.strip())
            section_type = None  # Reset for the next section
    doc_directory = os.path.join(settings.MEDIA_ROOT, 'documents')
    if not os.path.exists(doc_directory):
        os.makedirs(doc_directory)
    
    # Construct a unique filename for the document
    from datetime import datetime
    filename = f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc_path = os.path.join(doc_directory, filename)
    
    # Save the document
    document.save(doc_path)
    
    # Return the relative path of the document (from MEDIA_ROOT) for easy access
    return os.path.join('documents', filename)

def empty_documents_directory():
    documents_path = os.path.join(settings.MEDIA_ROOT, 'documents')
    for filename in os.listdir(documents_path):
        file_path = os.path.join(documents_path, filename)
        try:
            if os.path.isfile(file_path) or os.path.islink(file_path):
                os.unlink(file_path)
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
        except Exception as e:
            print(f'Failed to delete {file_path}. Reason: {e}')

def index(request):
    empty_documents_directory()
    if request.method == 'POST':
        raw_text = request.POST.get('text', '')
        structured_text = preprocessAI(raw_text)
        doc_rel_path = create_document_from_outline(structured_text)
        doc_abs_path = os.path.join(settings.MEDIA_ROOT, doc_rel_path)

        return render(request, 'home/index.html', {
            'structured_text': structured_text, 
            'doc_path': doc_abs_path
        })
    else:
        return render(request, 'home/index.html')

def download_document(request, doc_path):
    # Make sure that doc_path is the absolute path to the file
    file_path = os.path.join(settings.MEDIA_ROOT, doc_path)

    # Ensure the file exists
    if os.path.exists(file_path):
        with open(file_path, 'rb') as doc:
            response = HttpResponse(doc.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename="{}"'.format(os.path.basename(file_path))
        
        os.remove(file_path)
        return response
    else:
        return redirect('index')  # Redirect back to the form if no path is found